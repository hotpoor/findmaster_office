#!/bin/env python
#coding=utf-8
import sys
import os
import docx
from docx.shared import Length, Pt, RGBColor
from lxml import etree
import xmltodict
import json
import time
import string
import random
import requests
import jieba
import os
import subprocess

import qiniu

from setting import settings

from xml.dom.minidom import parseString

# file=docx.Document("template.docx")
# file=docx.Document("矿井安全监测监控管理机构及其职责2022.2.10.docx")
# file=docx.Document("彭庄煤矿2022年防灭火设计3.16.docx")
# filename = "4306轨道运输巷启封密闭恢复通风安全技术措施.docx"
# filename = "彭庄煤矿2022年防灭火设计3.16.docx"
# filename = "矿井安全监测监控管理机构及其职责2022.2.10.docx"
sub = 0
def postprocessor(path, key, value):
    global sub
    value = {
        "value":value,
        "key":key,
        "num":sub,
        "path":path,
        "content_id":"".join(random.choice(string.ascii_lowercase+string.digits) for _ in range(6))
    }
    key = "%s"%(sub)
    sub +=1
    return key, value

def postprocessor1(path, key, value):
    global sub
    value = {
        "value":value,
        "key":key,
        "num":sub,
        "path":path,
    }
    sub +=1
    return key, value

def file_md5(file_path):
    with open(file_path, 'rb') as fp:
        data = fp.read()
    file_md5= hashlib.md5(data).hexdigest()
    return file_md5

def upload_docx(filename,download_link="",save_block_id=None):
    sub = 0
    if filename.split(".")[-1] in ["docx"]:
        file=docx.Document(filename)
        body_xml_str = file._body._element.xml # 获取body中的xml
        body_xml = etree.fromstring(body_xml_str) # 转换成lxml结点
        # print(body_xml_str)
        # print(etree.tounicode(body_xml)) # 打印查看
        # xmltodict.parse()方法可以将xml数据转为python中的dict字典数据
        # parser_data = xmltodict.parse(body_xml_str)
        sub = 0
        parser_data = xmltodict.parse(etree.tounicode(body_xml),postprocessor=postprocessor)
        json_conversion1 = json.dumps(parser_data) 
        # print(json_conversion1)
        # with open("xml1.json", "w", encoding="utf-8") as bfile:
        #     bfile.write(json.dumps(json_conversion1, indent=2, ensure_ascii=False)) #ensure_ascii=False可以消除json包含中文的乱码问题

        sub = 0
        parser_data = xmltodict.parse(etree.tounicode(body_xml),postprocessor=postprocessor1)
        json_conversion2 = json.dumps(parser_data) 
        # print(json_conversion2)

        # with open("xml2.json", "w", encoding="utf-8") as bfile:
        #     bfile.write(json.dumps(json_conversion2, indent=2, ensure_ascii=False)) #ensure_ascii=False可以消除json包含中文的乱码问题

        dest = "./dest"
        file_docx_to_pdf = filename
        output = subprocess.check_output([settings["libreoffice"],"--headless","--convert-to","pdf",file_docx_to_pdf,"--outdir",dest])
        dest_file = "./dest/%s.pdf"%(file_docx_to_pdf.split("/")[-1].split(".docx")[0])


        # path_now = "%s/%s/%s"%(block_id,root,file_item)
        # path_current = "%s/%s"%(root,file_item)
        # file_md5_str = file_md5(path_current)
        file_md5_str = file_md5(dest_file)

        # file_item_list = file_item.split(".")
        qiniu_file_key = "%s_s_%s"%(save_block_id,file_md5_str)
        file_info = {
            # "name":file_item,
            # "path":path_now,
            "md5":file_md5_str,
            "link":"http://office-cdn-1.xialiwei.com/%s"%qiniu_file_key,
            # "cdn_type":"secret",
            "cdn_type":"public",
        }
        # files_list.append(file_info)
        # files_data.append(file_info)
        print(file_info)
        q = qiniu.Auth(settings["QiniuAccessKey"], settings["QiniuSecretKey"])
        # uptoken = q.upload_token("audio", key, 3600, policy)
        uptoken = q.upload_token("ofcourse-office-1")
        # localfile = os.path.join(os.path.dirname(__file__), path_current)
        # ret, info = qiniu.put_file(uptoken, qiniu_file_key, localfile)
        ret, info = qiniu.put_file(uptoken, qiniu_file_key, dest_file)

        pdf_download_link = file_info["link"]


        p_list = [
            "文件名: %s"%filename,
            "在线访问: <a target=\"_blank\" href=\"%s\">%s</a>"%(pdf_download_link,filename.split("/")[-1].replace(".docx",".pdf")),
            "下载地址: <a target=\"_blank\" href=\"%s?attname=%s\">%s</a>"%(download_link,filename.split("/")[-1],download_link)
        ]
        doc_dd_0 = json.loads(json_conversion1)
        for k,v in doc_dd_0.items():
            is_number = False
            try:
                k = int(k)
                is_number = True
            except Exception as e:
                is_number = False
            if is_number:
                for _k,_v in v["value"].items():
                    ais_number = False
                    try:
                        _k = int(_k)
                        ais_number = True
                    except Exception as e:
                        ais_number = False
                    if ais_number:
                        if _v["key"] in ["w:p","w:tbl","w:sectPr"]:
                            p = ""
                            p_add = False
                            for _kk,_vv in _v["value"].items():
                                if isinstance(_vv,dict):
                                    if _vv["key"] in ["w:r"]:
                                        if isinstance(_vv["value"],dict):
                                            for _kkk,_vvv in _vv["value"].items():
                                                if isinstance(_vvv,dict):
                                                    if _vvv["key"] in ["w:t"] and isinstance(_vvv["value"],str):
                                                        p = "%s%s"%(p,_vvv["value"])
                                                        p_add = True
                            if p_add:
                                p_list.append(p)
        p_list_str = json.dumps(p_list)
        print("文档行数:",len(p_list))
        url = "%s/api/page/add_free_docx"%(settings["host"])
        header = {
            'user-agent': 'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/86.0.4240.198 Safari/537.36'
        }
        data = {
            "token": "xialiwei_follows_god",
            "user_id": settings["user_id"],
            "title": filename,
            "desc": filename,
            "xml1":json_conversion1,
            "xml2":json_conversion2,
            "pre_p_list":p_list_str,
        }
        request = requests.post(url, headers=header, data=data)
        request_json = json.loads(request.text)
        block_id = request_json["block_id"]
        print(block_id)

        url = "%s/api/search/add_free_page"%(settings["host"])
        header = {
            'user-agent': 'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/86.0.4240.198 Safari/537.36'
        }
        data = {
            "token": "xialiwei_follows_god",
            "user_id": settings["user_id"],
            "search":"add free page default",
            "type":"sequence",
            "block_id":block_id,
            "block_id_sequence":"add free page default",
            "search":"add free page",
        }
        request = requests.post(url, headers=header, data=data)
        request_json = json.loads(request.text)

def upload_doc(filename,download_link="",save_block_id=None):
    dest = "./dest"
    file = filename
    output = subprocess.check_output([settings["libreoffice"],"--headless","--convert-to","docx",file,"--outdir",dest])
    print(output)
    dest_file = "./dest/%sx"%(file.split("/")[-1])
    print(dest_file)
    upload_docx(dest_file,download_link,save_block_id)
# print("=====")

# a = parseString(body_xml_str)
# all_data = []
# def get_child_nodes(item,data=[]):
#     if len(item.childNodes) == 0:
#         if hasattr(item,"nodeValue"):
#             data_str = item.nodeValue
#         else:
#             data_str = ""
#         info = {
#             "tag_name":item.nodeName,
#             "data":data_str,
#             "has_child":"false",
#             "childNodes":[]
#         }
#     else:
#         info = {
#             "tag_name":item.nodeName,
#             "data":"",
#             "has_child":"true",
#             "childNodes":[]
#         }
#         for i in item.childNodes:
#             j = get_child_nodes(i,[])
#             info["childNodes"].append(j)
#     data.append(info)
#     return data
# all_data = get_child_nodes(a)
# all_data_json = {"data":all_data}
# # print(all_data_json)
# with open("xml.json", "w", encoding="utf-8") as afile:
#     afile.write(json.dumps(all_data_json, indent=2, ensure_ascii=False)) #ensure_ascii=False可以消除json包含中文的乱码问题
